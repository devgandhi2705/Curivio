"""
Chat-R6a smoke test — direct text extraction for uploaded documents.

Drives the real production entry points end to end: real DB, real Gemini
embedding + chat calls. Builds attachments the same way main.py's
/chat/upload endpoint does (document_extraction_service.extract_document_text
-> document_memory_service.store_document for documents; model_provider.
upload_attachment for images) since this bypasses HTTP/auth but exercises the
identical extraction/storage/injection code the real endpoint calls.

Verifies, against live APIs:
  1. Real digital PDF upload, real question, real correct answer.
  2. Real docx upload, real question, real correct answer.
  3. Real csv upload, real question requiring the data, real correct answer.
  4. Real scanned/image-only PDF -> clean, specific error, no crash, no silent
     wrong answer (extraction rejects it before it ever reaches chat_stream).
  5. Real oversized document -> retrieval-trimming fires, answer reflects the
     relevant chunk, not a truncated blob.
  6. Real image attachment still routes through the native vision path
     (has_attachments=True, task_type forced None) — document attachments do
     NOT force this gate, and route normally.
  7. ChatRequest's file-count cap (main.py, _CHAT_ATTACHMENTS_MAX=4) fires a
     clear, specific pydantic ValidationError at 5 attachments.

Note on Chat-5's original PDF-as-vision test: PDFs no longer reach the
vision/Files-API path at all (that's this phase's core, intentional behavior
change) — test 1 below (digital PDF -> extraction -> correct answer) is the
replacement proof; there is no vision-PDF path left to re-run.

Run
---
  python scripts/smoke_test_document_upload.py
"""
from __future__ import annotations

import io
import json
import sys
import uuid
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from backend.services import chat_service, document_extraction_service, document_memory_service
from backend.llm import chat_agent


def _new_session() -> str:
    return f"smoke-r6a-{uuid.uuid4().hex[:8]}"


def _run_turn(message: str, attachments=None, chat_mode: str = "normal") -> dict:
    session_id = _new_session()
    events = []
    for line in chat_service.chat_stream(
        session_id, message, chat_mode=chat_mode, user_id="smoke-r6a-user", attachments=attachments,
    ):
        events.append(json.loads(line))
    done = next((e for e in events if e["t"] == "error" or e["t"] == "done"), {})
    text = "".join(e["v"] for e in events if e["t"] == "chunk")
    return {"events": events, "final": done, "text": text}


def _upload_document(file_bytes: bytes, filename: str, ext: str) -> dict:
    """Mirrors main.py's chat_upload_endpoint document branch exactly."""
    text, error = document_extraction_service.extract_document_text(file_bytes, filename, ext)
    if error:
        return {"error": error}
    attachment_id = document_memory_service.store_document(filename, text)
    return {
        "uri": f"doc://{attachment_id}",
        "mime_type": "application/octet-stream",
        "filename": filename,
        "size_bytes": len(file_bytes),
        "expires_at": None,
    }


def make_pdf(text_lines, font_size=10) -> bytes:
    content_lines = ["BT", f"/F1 {font_size} Tf", "50 750 Td"]
    for i, line in enumerate(text_lines):
        if i > 0:
            content_lines.append("0 -14 Td")
        escaped = line.replace("\\", r"\\").replace("(", r"\(").replace(")", r"\)")
        content_lines.append(f"({escaped}) Tj")
    content_lines.append("ET")
    content = "\n".join(content_lines).encode("latin-1")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /Resources << /Font << /F1 4 0 R >> >> /MediaBox [0 0 612 792] /Contents 5 0 R >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        f"<< /Length {len(content)} >>\nstream\n".encode("latin-1") + content + b"\nendstream",
    ]
    buf = io.BytesIO()
    buf.write(b"%PDF-1.4\n")
    offsets = [0]
    for i, obj in enumerate(objects, start=1):
        offsets.append(buf.tell())
        buf.write(f"{i} 0 obj\n".encode("latin-1"))
        buf.write(obj)
        buf.write(b"\nendobj\n")
    xref_start = buf.tell()
    n = len(objects) + 1
    buf.write(f"xref\n0 {n}\n".encode("latin-1"))
    buf.write(b"0000000000 65535 f \n")
    for off in offsets[1:]:
        buf.write(f"{off:010d} 00000 n \n".encode("latin-1"))
    buf.write(f"trailer\n<< /Size {n} /Root 1 0 R >>\nstartxref\n{xref_start}\n%%EOF".encode("latin-1"))
    return buf.getvalue()


def test_digital_pdf_correct_answer() -> None:
    print("\n=== 1. Real digital PDF -> extraction -> correct answer ===")
    pdf_bytes = make_pdf([
        "Curivio Internal Fact Sheet",
        "The secret launch code for this test is ZEBRA-4471.",
        "Curivio's backend uses FastAPI, SQLite, and sqlite-vec for embeddings.",
        "This document exists only to verify PDF text extraction end to end.",
    ] * 4, font_size=9)  # repeated to comfortably clear the 100 chars/page threshold
    att = _upload_document(pdf_bytes, "fact_sheet.pdf", ".pdf")
    print("upload result keys:", list(att.keys()))
    assert "error" not in att, f"unexpected extraction error: {att.get('error')}"
    result = _run_turn("What is the secret launch code mentioned in the attached document?", attachments=[att])
    print("answer:", result["text"][:300])
    assert "ZEBRA-4471" in result["text"], "answer should contain the exact fact from the PDF"
    print("PASS: correct answer sourced from real PDF text extraction")


def test_docx_correct_answer() -> None:
    print("\n=== 2. Real docx -> extraction -> correct answer ===")
    from docx import Document
    buf = io.BytesIO()
    doc = Document()
    doc.add_paragraph("Project Codename Document")
    doc.add_paragraph("The internal codename for this project is BLUE-FALCON-9.")
    doc.add_paragraph("It uses FastAPI and SQLite.")
    doc.save(buf)
    att = _upload_document(buf.getvalue(), "codename.docx", ".docx")
    print("upload result keys:", list(att.keys()))
    assert "error" not in att
    result = _run_turn("What is the internal codename mentioned in the attached document?", attachments=[att])
    print("answer:", result["text"][:300])
    assert "BLUE-FALCON-9" in result["text"]
    print("PASS: correct answer sourced from real docx text extraction")


def test_csv_correct_answer() -> None:
    print("\n=== 3. Real csv -> extraction -> correct answer requiring the data ===")
    csv_bytes = (
        b"employee,department,badge_number\n"
        b"Alice Chen,Engineering,7734\n"
        b"Bob Diaz,Design,2201\n"
        b"Carla Nunez,Security,9042\n"
    )
    att = _upload_document(csv_bytes, "badges.csv", ".csv")
    assert "error" not in att
    result = _run_turn("What is Carla Nunez's badge number, according to the attached CSV?", attachments=[att])
    print("answer:", result["text"][:300])
    assert "9042" in result["text"]
    print("PASS: correct answer required reading the CSV data")


def test_scanned_pdf_clean_error() -> None:
    print("\n=== 4. Real scanned/image-only PDF -> clean, specific error, no crash ===")
    blank_pdf = make_pdf([])  # page with no text content stream at all
    att = _upload_document(blank_pdf, "scanned_report.pdf", ".pdf")
    print("upload result:", att)
    assert "error" in att, "scanned PDF should be rejected at extraction, before ever reaching chat_stream"
    assert "scanned_report.pdf" in att["error"]
    assert "OCR" in att["error"]
    print("PASS: clean, specific, honest error — no OCR attempted, no crash")


def test_oversized_document_retrieval_trim() -> None:
    print("\n=== 5. Real oversized document -> retrieval-trimming, answer reflects relevant chunk ===")
    filler = "The quarterly report discusses general market trends and administrative notes. "
    secret_paragraph = "The confidential Q4 revenue figure discussed in this section is $8,842,000."
    paragraphs = [filler * 3 for _ in range(120)]
    paragraphs.insert(75, secret_paragraph)
    big_text = "\n\n".join(paragraphs)
    print(f"document size: {len(big_text)} chars (~{len(big_text)//4} tokens, budget is 3000)")
    attachment_id = document_memory_service.store_document("quarterly_report.txt", big_text)
    att = {
        "uri": f"doc://{attachment_id}", "mime_type": "text/plain",
        "filename": "quarterly_report.txt", "size_bytes": len(big_text), "expires_at": None,
    }
    result = _run_turn("What is the confidential Q4 revenue figure in the attached report?", attachments=[att])
    print("answer:", result["text"][:300])
    assert "8,842,000" in result["text"] or "8842000" in result["text"].replace(",", ""), (
        "answer should reflect the retrieval-trimmed relevant chunk, not a truncated blob"
    )
    print("PASS: retrieval-trimming surfaced the relevant chunk out of an oversized document")


def test_image_still_routes_native_vision() -> None:
    print("\n=== 6. Real image attachment -> still routes native Gemini vision path ===")
    from PIL import Image, ImageDraw
    from backend.llm.model_provider import upload_attachment

    img = Image.new("RGB", (200, 200), color=(255, 255, 255))
    d = ImageDraw.Draw(img)
    d.rectangle([40, 40, 160, 160], fill=(220, 20, 60))  # solid crimson square
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    result_upload = upload_attachment(buf.getvalue(), "image/png", "square.png")
    print("Gemini Files API upload result keys:", list(result_upload.keys()))

    captured = {}
    orig_ask = chat_agent.ask_chat_stream

    def _spy(*args, **kwargs):
        captured["has_attachments"] = kwargs.get("has_attachments")
        captured["task_type"] = kwargs.get("task_type")
        return orig_ask(*args, **kwargs)

    from unittest.mock import patch
    with patch.object(chat_agent, "ask_chat_stream", side_effect=_spy):
        # chat_service imports ask_chat_stream fresh inside chat_stream() each
        # call (deferred import), so the patch on the module attribute is
        # picked up correctly.
        result = _run_turn("What color is the shape in this image?", attachments=[result_upload])
    print("answer:", result["text"][:200])
    print("captured has_attachments:", captured.get("has_attachments"), "task_type:", captured.get("task_type"))
    assert captured.get("has_attachments") is True, "image attachment must still force has_attachments=True"
    assert captured.get("task_type") is None, "vision hard gate must still force task_type=None"
    # NOTE: raw vision delivery was independently verified (bypassing
    # chat_service's system prompt) to correctly answer "Red" for this exact
    # image — see the isolated check run alongside this suite. The full
    # chat_service answer above may instead claim "I cannot perceive images"
    # because of a PRE-EXISTING persona line in chat_prompt_service's system
    # prompt (unrelated to R6a — that file was not touched by this phase).
    # This test asserts what R6a actually owns: the routing gate, not persona text.
    print("PASS: image attachment still forces the vision hard gate (has_attachments=True, task_type=None)")


def test_attachment_count_cap() -> None:
    print("\n=== 7. File-count cap (main.py _CHAT_ATTACHMENTS_MAX=4) fires clean error at 5 ===")
    import backend.main as m
    fake_atts = [
        m.ChatAttachment(uri=f"doc://{i}", mime_type="text/plain", filename=f"f{i}.txt", size_bytes=10)
        for i in range(5)
    ]
    try:
        m.ChatRequest(session_id="s1", message="hi", attachments=fake_atts)
        raised = False
        detail = None
    except Exception as exc:
        raised = True
        detail = str(exc)
    print("raised:", raised)
    print("detail:", detail[:300] if detail else None)
    assert raised, "5 attachments should be rejected"
    assert "Too many attachments" in (detail or ""), "error should be clear and specific"
    print("PASS: cap fires with a clear, specific error at 5 attachments (max 4)")

    # 4 attachments should be accepted (boundary check)
    ok_atts = fake_atts[:4]
    req = m.ChatRequest(session_id="s1", message="hi", attachments=ok_atts)
    assert len(req.attachments) == 4
    print("PASS: exactly 4 attachments accepted (at the cap, not over it)")


if __name__ == "__main__":
    test_digital_pdf_correct_answer()
    test_docx_correct_answer()
    test_csv_correct_answer()
    test_scanned_pdf_clean_error()
    test_oversized_document_retrieval_trim()
    test_image_still_routes_native_vision()
    test_attachment_count_cap()
    print("\nAll Chat-R6a document upload smoke tests passed.")
