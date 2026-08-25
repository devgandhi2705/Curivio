"""
Phase P — document_extraction_service now returns (text, pages, error, error_type),
a stable machine code per distinct failure reason instead of just a message
string. These tests prove each reachable reason is real (built from real
pypdf/docx/decode behavior, not mocked) and distinguishable from the others.

pages (citation grounding phase): real per-page text for PDFs, None for
docx/plain-text (no real page concept -- see TestPdfSuccess/TestDocxReasons/
TestPlainTextReasons below).

ERROR_TEXT_UNDECODABLE is not exercised here: _extract_plain's fallback is
utf-8 -> latin-1, and latin-1 maps every byte value 0-255 to a character, so
that except branch is realistically unreachable with real bytes (pre-existing,
not touched by this phase).
"""
from __future__ import annotations

import io
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from docx import Document as DocxDocument
from pypdf import PdfWriter
from reportlab.pdfgen import canvas

from backend.services import document_extraction_service as des


def _real_text_pdf_bytes() -> bytes:
    buf = io.BytesIO()
    c = canvas.Canvas(buf)
    c.drawString(100, 750, "Hello world, this is a real extractable PDF text line.")
    c.drawString(100, 730, "A second line so the page clears the 100 chars/page floor.")
    c.save()
    return buf.getvalue()


def _blank_pdf_bytes() -> bytes:
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


def _zero_page_pdf_bytes() -> bytes:
    writer = PdfWriter()
    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


def _real_text_docx_bytes() -> bytes:
    doc = DocxDocument()
    doc.add_paragraph("Real extractable docx paragraph.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _empty_docx_bytes() -> bytes:
    doc = DocxDocument()
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestPdfSuccess:
    def test_real_pdf_extracts_text_no_error(self):
        text, pages, error, error_type = des.extract_document_text(_real_text_pdf_bytes(), "r.pdf", ".pdf")
        assert error is None
        assert error_type is None
        assert "Hello world" in text
        # Citation grounding: PDFs get real per-page text, 1-based page_no.
        assert pages is not None
        assert pages[0]["page_no"] == 1
        assert "Hello world" in pages[0]["text"]


class TestPdfFailureReasons:
    def test_corrupt_pdf_bytes(self):
        text, pages, error, error_type = des.extract_document_text(b"not a pdf at all", "bad.pdf", ".pdf")
        assert text is None
        assert pages is None
        assert error_type == des.ERROR_PDF_UNREADABLE
        assert "bad.pdf" in error

    def test_zero_page_pdf(self):
        text, pages, error, error_type = des.extract_document_text(_zero_page_pdf_bytes(), "empty.pdf", ".pdf")
        assert text is None
        assert pages is None
        assert error_type == des.ERROR_PDF_NO_PAGES

    def test_scanned_image_only_pdf(self):
        text, pages, error, error_type = des.extract_document_text(_blank_pdf_bytes(), "scan.pdf", ".pdf")
        assert text is None
        assert pages is None
        assert error_type == des.ERROR_PDF_SCANNED
        assert "scanned" in error.lower()


class TestDocxReasons:
    def test_real_docx_extracts_text_no_error(self):
        text, pages, error, error_type = des.extract_document_text(_real_text_docx_bytes(), "r.docx", ".docx")
        assert error is None
        assert error_type is None
        assert "Real extractable docx paragraph." in text
        # docx has no real page concept -- must stay honestly NULL, never fabricated.
        assert pages is None

    def test_corrupt_docx_bytes(self):
        text, pages, error, error_type = des.extract_document_text(b"not a docx at all", "bad.docx", ".docx")
        assert text is None
        assert pages is None
        assert error_type == des.ERROR_DOCX_UNREADABLE

    def test_empty_docx_no_text(self):
        text, pages, error, error_type = des.extract_document_text(_empty_docx_bytes(), "blank.docx", ".docx")
        assert text is None
        assert pages is None
        assert error_type == des.ERROR_DOCX_NO_TEXT


class TestPlainTextReasons:
    def test_real_text_extracts_no_error(self):
        text, pages, error, error_type = des.extract_document_text(b"real content here", "r.txt", ".txt")
        assert error is None
        assert error_type is None
        assert text == "real content here"
        # plain text/code has no real page concept -- must stay honestly NULL.
        assert pages is None

    def test_whitespace_only_text_is_empty(self):
        text, pages, error, error_type = des.extract_document_text(b"   \n\t  ", "blank.txt", ".txt")
        assert text is None
        assert pages is None
        assert error_type == des.ERROR_TEXT_EMPTY


class TestUnsupportedExtension:
    def test_unknown_extension(self):
        text, pages, error, error_type = des.extract_document_text(b"whatever", "f.xyz", ".xyz")
        assert text is None
        assert pages is None
        assert error_type == des.ERROR_UNSUPPORTED_EXTENSION
        assert ".xyz" in error


class TestErrorTypesAreAllDistinct:
    def test_all_eight_constants_are_unique_strings(self):
        codes = {
            des.ERROR_PDF_UNREADABLE, des.ERROR_PDF_NO_PAGES, des.ERROR_PDF_SCANNED,
            des.ERROR_DOCX_UNREADABLE, des.ERROR_DOCX_NO_TEXT,
            des.ERROR_TEXT_UNDECODABLE, des.ERROR_TEXT_EMPTY, des.ERROR_UNSUPPORTED_EXTENSION,
        }
        assert len(codes) == 8


if __name__ == "__main__":
    import pytest
    sys.exit(pytest.main([__file__, "-v"]))
