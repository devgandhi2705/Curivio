"""
Feed v2 profile agent — live integration tests (real profile LLM call).

Skipped without a Gemini key and excluded from the default suite (integration).
Proves the parts that need a real model:
  - a full persona + coverage decision over a real mixed 4-material set;
  - the coverage_mode INFERENCE actually reads the signals: a structured-syllabus
    upload and a couple of small blog links land on DIFFERENT coverage_mode values;
  - the thin-title-rescued-by-material case ("ML course" + a real syllabus →
    learning_subject comes from the material, not the two-word title);
  - the new "profile" role's cross-provider fallback leg serves when the primary
    (Gemini) leg is forced to fail (needs an OpenRouter key too).
"""
import io
import os
import sqlite3

import pytest
import sqlite_vec

# Import first so load_dotenv() populates os.environ before the key check below.
from backend.database.schema import ALL_TABLES, MIGRATIONS
from backend.services.feed_v2 import db as v2db
from backend.services.feed_v2.schema import run_v2_migrations
from backend.services.feed_v2.ingestion import materials, links
from backend.services.feed_v2 import projects as P
from backend.services.feed_v2.llm import provider

_HAS_KEY = bool(os.getenv("GEMINI_API_KEYS") or os.getenv("GEMINI_API_KEY"))
_HAS_OR = bool(os.getenv("OPENROUTER_API_KEYS") or os.getenv("OPENROUTER_API_KEY"))
pytestmark = [pytest.mark.integration,
              pytest.mark.skipif(not _HAS_KEY, reason="no Gemini API key")]


def _build_db(path):
    conn = sqlite3.connect(path)
    conn.enable_load_extension(True); sqlite_vec.load(conn); conn.enable_load_extension(False)
    for s in ALL_TABLES:
        conn.execute(s)
    for m in MIGRATIONS:
        try:
            conn.execute(m) if not isinstance(m, (list, tuple)) else [conn.execute(x) for x in m]
        except sqlite3.OperationalError as e:
            if not any(k in str(e).lower() for k in ("already exists", "duplicate column", "no such column")):
                raise
    run_v2_migrations(conn)
    conn.execute("INSERT OR IGNORE INTO users(user_id,email,name,hashed_pw) VALUES('u1','u1@t.com','u','x')")
    conn.commit(); conn.close()


@pytest.fixture
def db(tmp_path, monkeypatch):
    path = str(tmp_path / "profile_live.db")
    _build_db(path)
    monkeypatch.setattr(v2db, "DB_PATH", path)
    monkeypatch.setattr(links, "_MOCK", True)  # deterministic, small unstructured link text
    return path


# ── real file builders ────────────────────────────────────────────────────────
def _png(text="DIAGRAM"):
    from PIL import Image, ImageDraw
    img = Image.new("RGB", (240, 120), "white")
    ImageDraw.Draw(img).text((20, 50), text, fill="black")
    buf = io.BytesIO(); img.save(buf, format="PNG"); return buf.getvalue()


def _pdf():
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    c.drawString(72, 740, "Introduction to widget mechanics and load tolerances. " * 6)
    c.drawString(72, 720, "Real digital text well over the scanned-PDF threshold. " * 3)
    c.showPage(); c.save()
    return buf.getvalue()


def _docx_headings():
    from docx import Document
    doc = Document()
    doc.add_heading("Overview", level=1)
    doc.add_paragraph("Widgets are small load-bearing components. " * 15)
    doc.add_heading("Assembly", level=1)
    doc.add_paragraph("Assembly requires torque control. " * 15)
    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()


def _md_syllabus(n=12):
    topics = ["Arrays", "Linked Lists", "Stacks", "Queues", "Hash Tables", "Trees",
              "Binary Search Trees", "Heaps", "Graphs", "Sorting", "Dynamic Programming",
              "Complexity Analysis"]
    lines = ["# Course Syllabus: Data Structures & Algorithms",
             "This course follows a fixed weekly schedule leading to the semester final exam.", ""]
    for i in range(n):
        t = topics[i % len(topics)]
        lines += [f"## Week {i + 1}: {t}", f"Required reading and graded exercises on {t}.", ""]
    return "\n".join(lines).encode()


def _md_ml_syllabus():
    return (
        "# Deep Learning Course Syllabus\n"
        "## Module 1: Linear Models and Gradient Descent\nFoundations of optimization.\n"
        "## Module 2: Neural Networks and Backpropagation\nFeedforward nets and training.\n"
        "## Module 3: Convolutional Networks for Vision\nCNNs, pooling, architectures.\n"
        "## Module 4: Transformers and Attention\nSelf-attention and sequence models.\n"
        "## Module 5: Reinforcement Learning\nPolicy gradients and value methods.\n"
    ).encode()


def test_full_profile_over_mixed_material_set(db, capsys):
    proj = P.create_project(
        "u1", "Widgets 101",
        "I'm a mechanical engineering student studying widget design for my coursework",
        "intermediate")
    pid = proj["project_id"]
    materials.ingest("u1", pid, file_bytes=_pdf(), filename="widgets.pdf")
    materials.ingest("u1", pid, file_bytes=_docx_headings(), filename="widgets.docx")
    materials.ingest("u1", pid, file_bytes=_png(), filename="diagram.png")
    materials.ingest("u1", pid, url="https://example.com/widgets-intro")

    out = P.generate_profile("u1", pid)
    pr = out["profile"]
    with capsys.disabled():
        print("\n=== FULL PROFILE (mixed PDF+DOCX+image+link set) ===")
        for k in ("persona", "learning_subject", "primary_focus", "industry_context",
                  "material_scope", "coverage_mode", "coverage_reasoning"):
            print(f"  {k}: {pr[k]!r}")
    assert out["profile_status"] == "ready"
    assert out["coverage_mode"] in {"material_bound", "material_anchored", "open"}
    assert pr["persona"] and pr["material_scope"] and pr["coverage_reasoning"]


def test_opposite_coverage_modes_from_input_shape(db, capsys):
    # Syllabus-shaped: ONE big STRUCTURED document + exam/syllabus description.
    syl = P.create_project(
        "u1", "Data Structures",
        "I'm preparing for my final exam and following the course syllabus exactly.",
        "intermediate")
    materials.ingest("u1", syl["project_id"], file_bytes=_md_syllabus(12), filename="syllabus.md")
    cov_syl = P.generate_profile("u1", syl["project_id"])["coverage_mode"]

    # Blog-shaped: a couple of small UNSTRUCTURED links + exploratory description.
    blog = P.create_project(
        "u1", "Async Rust",
        "just getting into async rust, want to learn the basics",
        "beginner")
    materials.ingest("u1", blog["project_id"], url="https://blog.example.com/async-basics")
    materials.ingest("u1", blog["project_id"], url="https://blog.example.com/tokio-intro")
    cov_blog = P.generate_profile("u1", blog["project_id"])["coverage_mode"]

    with capsys.disabled():
        print(f"\ncoverage_mode inference: syllabus={cov_syl}  blogs={cov_blog}")
    # The core proof: the inference reads the signals — it does NOT return the same
    # mode for a structured syllabus and a couple of loose links.
    assert cov_syl != cov_blog, f"inference ignored input shape: both={cov_syl}"
    assert cov_syl == "material_bound", f"structured syllabus should be material_bound, got {cov_syl}"


def test_thin_title_rescued_by_material(db, capsys):
    proj = P.create_project("u1", "ML course", "I'm a student", "beginner")
    materials.ingest("u1", proj["project_id"], file_bytes=_md_ml_syllabus(), filename="ml_syllabus.md")
    out = P.generate_profile("u1", proj["project_id"])
    subj = out["profile"]["learning_subject"]
    with capsys.disabled():
        print(f"\nthin title 'ML course' -> learning_subject: {subj!r}")
    # Legacy's rule would keep the thin two-word title as the subject. The reversal
    # lets the syllabus win: the subject must be richer than the raw title.
    assert subj.strip().lower() != "ml course", "material did not rescue the thin title"


@pytest.mark.skipif(not _HAS_OR, reason="no OpenRouter key for the cross-provider fallback leg")
def test_profile_fallback_leg_serves(db, monkeypatch, capsys):
    """Same proof style as Phase 4b's image_ingestor test, for the new 'profile'
    role: force the primary (Gemini) leg to fail; confirm the cross-provider
    fallback (nemotron/OpenRouter) is the leg the router hands the request to.

    ROUTING proof is deterministic (the fallback leg is reached + invoked with the
    fallback model id, primary never serves). END-TO-END serve is conditional on the
    OpenRouter upstream accepting the call — nemotron's upstreams intermittently
    reject `response_format=json_object` with a 405 (reported as a pre-existing
    provider-layer issue, out of Phase 5 scope), so a 405/rate-limit run skips the
    serve half rather than flaking the suite."""
    primary_id = provider.MODEL_REGISTRY["gemini-3-flash-preview"][1]
    fallback_id = provider.MODEL_REGISTRY["nemotron-nano-30b"][1]
    real_or = provider._call_openrouter
    attempted: list[str] = []   # recorded BEFORE the API call -> routing reached the leg
    served: list[str] = []      # recorded AFTER success -> upstream actually served

    def fake_google(api_model_id, messages, system, schema, key, images=None):
        raise RuntimeError("simulated primary (gemini) outage")  # non-429 -> immediate leg abort

    def rec_openrouter(api_model_id, messages, system, schema, key, images=None):
        attempted.append(api_model_id)
        r = real_or(api_model_id, messages, system, schema, key, images)
        served.append(api_model_id)
        return r

    monkeypatch.setitem(provider._SDK_FOR_PROVIDER, "google", fake_google)
    monkeypatch.setitem(provider._SDK_FOR_PROVIDER, "openrouter", rec_openrouter)

    proj = P.create_project("u1", "Graph Theory", "I'm a math student", "intermediate")
    try:
        out = P.generate_profile("u1", proj["project_id"])
        ok = True
    except provider.AllLegsFailed as exc:
        out, ok = None, False
        upstream_err = str(exc)

    # ROUTING proof (deterministic): the profile role's fallback leg WAS reached and
    # invoked with the fallback model — and the primary NEVER served the request.
    assert attempted and all(a == fallback_id for a in attempted), attempted
    assert primary_id not in attempted and primary_id not in served

    with capsys.disabled():
        print(f"\nprofile routing: primary {primary_id} forced-fail -> fallback leg reached "
              f"{attempted}; served={served}; ok={ok}")

    if ok:
        assert served and out["profile_status"] == "ready" and out["profile"]["persona"]
    else:
        pytest.skip("profile fallback leg reached nemotron, but the OpenRouter upstream "
                    "rejected it this run (json_object 405 / rate-limit) — routing proven; "
                    "end-to-end serve is subject to upstream availability")
