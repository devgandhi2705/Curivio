"""
Feed v2 ingestion — offline tests (no network, run in the default suite).

Covers the paths that never reach chunking/embedding: structure extraction
(step 1/6 signal), and the corrupt-file failure-isolation guarantee (step 7) —
a failed material is a terminal 'failed' row that never wedges the project or
blocks the next material. The live chunk/embed/figure/vision paths are in
tests/test_feed_v2_ingestion.py (integration).
"""
import io
import sqlite3

import pytest
from docx import Document

from backend.database.schema import ALL_TABLES, MIGRATIONS
from backend.services.feed_v2 import db as v2db
from backend.services.feed_v2.schema import run_v2_migrations
from backend.services.feed_v2.ingestion import documents, materials, figures, chunking


def _build_db(path):
    conn = sqlite3.connect(path)  # conftest patches connect to load sqlite_vec
    for stmt in ALL_TABLES:
        conn.execute(stmt)
    for m in MIGRATIONS:
        try:
            conn.execute(m) if not isinstance(m, (list, tuple)) else [conn.execute(s) for s in m]
        except sqlite3.OperationalError as e:
            if not any(p in str(e).lower() for p in ("already exists", "duplicate column", "no such column")):
                raise
    run_v2_migrations(conn)
    conn.execute("INSERT OR IGNORE INTO users(user_id,email,name,hashed_pw) VALUES('u1','u1@t.com','u','x')")
    # v2_materials.project_id FKs v2_projects — the real flow creates the project
    # before uploading materials to it, so the fixture does too.
    for pid in ("p1", "p2"):
        conn.execute("INSERT OR IGNORE INTO v2_projects(project_id,user_id,name) VALUES(?,'u1',?)", (pid, pid))
    conn.commit()
    conn.close()


@pytest.fixture
def temp_db(tmp_path, monkeypatch):
    path = str(tmp_path / "ingest.db")
    _build_db(path)
    monkeypatch.setattr(v2db, "DB_PATH", path)
    return path


# ── structure extraction (step 1/6 signal) — pure, no DB/network ──────────────
def _docx_bytes(headings):
    doc = Document()
    for h in headings:
        doc.add_heading(h, level=1)
        doc.add_paragraph("Body text under " + h + ". " * 5)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_docx_structure_extracted():
    res = documents.extract(_docx_bytes(["Chapter One", "Chapter Two", "Chapter Three"]),
                            "book.docx", ".docx")
    assert res.error is None
    assert res.has_structure and res.section_count == 3
    assert [s["title"] for s in res.structure] == ["Chapter One", "Chapter Two", "Chapter Three"]


def test_markdown_structure_extracted():
    md = b"# Title\n\nintro\n\n## Section A\n\ntext\n\n## Section B\n\nmore"
    res = documents.extract(md, "notes.md", ".md")
    assert res.has_structure and res.section_count == 3


def test_plain_text_has_no_structure():
    res = documents.extract(b"just some flat text, no headings at all", "flat.txt", ".txt")
    assert res.error is None
    assert not res.has_structure and res.section_count == 0


def test_corrupt_pdf_returns_error_not_crash():
    res = documents.extract(b"%PDF-1.4 this is not a real pdf body", "broken.pdf", ".pdf")
    assert res.text is None
    assert res.error is not None
    assert res.sha256  # hash still computed on the raw bytes


# ── failure isolation (step 7) — DB, no network (extraction fails first) ──────
def test_corrupt_material_writes_failed_row_and_does_not_block_next(temp_db):
    r1 = materials.ingest_document("u1", "p1", b"not a real pdf", "broken.pdf")
    assert r1["status"] == "failed" and r1["error"]

    # A second ingest after the failure still produces its own row — not wedged.
    r2 = materials.ingest(u := "u1", "p1", file_bytes=b"\x00\x01", filename="mystery.xyz")
    assert r2["status"] == "failed"  # unsupported type -> terminal failed row

    with v2db.get_connection() as conn:
        rows = conn.execute(
            "SELECT material_id, type, extraction_status, extraction_error FROM v2_materials "
            "WHERE project_id='p1' ORDER BY created_at"
        ).fetchall()
    assert len(rows) == 2
    assert all(r["extraction_status"] == "failed" for r in rows)
    assert all(r["extraction_error"] for r in rows)


# ── FLAG 2: file storage is a distinct outcome from text extraction ──────────
def _pdf_with_two_images():
    """Two pages, one embedded image each -> two extractable figures, so one can
    fail to store while the other succeeds."""
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.utils import ImageReader
    from PIL import Image
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    for i, color in enumerate(("blue", "green")):
        c.drawString(72, 740, "Digital text well over the scanned-PDF threshold. " * 4)
        c.drawImage(ImageReader(Image.new("RGB", (120, 120), color)), 72, 500, width=150, height=150)
        c.drawString(72, 480, f"Figure {i+1}: a {color} square")
        c.showPage()
    c.save()
    return buf.getvalue()


class _FakeR2:
    """boto3 client stand-in. put_object RAISES on the first call, succeeds after —
    so exactly one figure's real _r2_upload try/except converts it to a store
    failure while the other stores cleanly."""
    def __init__(self, fail_calls=(1,)):
        self.calls = 0
        self.fail_calls = set(fail_calls)
    def put_object(self, **kw):
        self.calls += 1
        if self.calls in self.fail_calls:
            raise RuntimeError("boto3 put_object simulated failure")
        return {}


def _mock_no_network(monkeypatch):
    # Keep the FLAG 2 tests offline+deterministic: no real embedding or chunking.
    monkeypatch.setattr(figures, "embed_query", lambda text: [0.0])
    monkeypatch.setattr(chunking, "chunk_and_embed",
                        lambda *a, **k: {"chunk_count": 0, "embedding_count": 0})
    monkeypatch.setattr(chunking, "chunk_and_embed_pages",   # Phase 8b: PDFs route here now
                        lambda *a, **k: {"chunk_count": 0, "embedding_count": 0})
    monkeypatch.setenv("R2_BUCKET_NAME", "test-bucket")  # get past _r2_upload's bucket check


def test_partial_figure_storage_failure_is_recorded_distinctly(temp_db, monkeypatch):
    _mock_no_network(monkeypatch)
    fake = _FakeR2(fail_calls=(1,))                       # first figure fails, second succeeds
    monkeypatch.setattr(figures, "_r2_client", lambda: fake)
    r = materials.ingest_document("u1", "p1", _pdf_with_two_images(), "doc.pdf")
    assert r["status"] == "done" and r["figure_count"] == 2

    with v2db.get_connection() as conn:
        row = conn.execute(
            "SELECT extraction_status, extraction_error, storage_status FROM v2_materials "
            "WHERE material_id=?", (r["material_id"],)).fetchone()
        keys = [f["image_key"] for f in conn.execute(
            "SELECT image_key FROM v2_material_figures WHERE material_id=? ORDER BY page_no",
            (r["material_id"],)).fetchall()]
    # Text fine, but one figure image did not store — distinct, not collapsed into 'done'.
    assert row["extraction_status"] == "done"
    assert row["extraction_error"] is None
    assert row["storage_status"] == "degraded"
    assert fake.calls == 2 and 1 in fake.fail_calls       # boto3 was hit twice, one raised
    assert keys.count(None) == 1 and sum(k is not None for k in keys) == 1  # one failed, one stored
    print("PARTIAL-FAILURE row:", dict(row), "| figure image_keys:", keys)


def test_full_success_row_unchanged_when_storage_ok(temp_db, monkeypatch):
    _mock_no_network(monkeypatch)
    fake = _FakeR2(fail_calls=())                         # every put_object succeeds
    monkeypatch.setattr(figures, "_r2_client", lambda: fake)
    r = materials.ingest_document("u1", "p1", _pdf_with_two_images(), "doc.pdf")
    with v2db.get_connection() as conn:
        row = conn.execute(
            "SELECT extraction_status, extraction_error, storage_status FROM v2_materials "
            "WHERE material_id=?", (r["material_id"],)).fetchone()
        keys = [f["image_key"] for f in conn.execute(
            "SELECT image_key FROM v2_material_figures WHERE material_id=?",
            (r["material_id"],)).fetchall()]
    assert row["extraction_status"] == "done"
    assert row["extraction_error"] is None
    assert row["storage_status"] == "ok"                  # full success: 'ok', clearly != 'degraded'
    assert all(k is not None for k in keys)               # every figure image stored
    print("FULL-SUCCESS row:", dict(row), "| figure image_keys:", keys)


def test_signal_columns_are_queryable(temp_db):
    # A failed material still lands with the queryable coverage columns populated.
    materials.ingest_document("u1", "p2", b"bad pdf bytes", "x.pdf")
    with v2db.get_connection() as conn:
        row = conn.execute(
            "SELECT type, has_structure, section_count FROM v2_materials WHERE project_id='p2'"
        ).fetchone()
    assert row["type"] == "document"
    assert row["has_structure"] == 0 and row["section_count"] == 0
