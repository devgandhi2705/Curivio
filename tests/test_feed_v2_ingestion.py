"""
Feed v2 ingestion — live integration tests (real Gemini embeddings + vision).

Skipped without a Gemini key and excluded from the default suite (marked
integration). Exercises the full pipeline end to end:
  - a mixed upload set (PDF w/ real embedded figure, DOCX, plain image, link)
    all land in v2_materials with correct type + status
  - a figure extracted from the PDF with a real page number + caption
  - chunk count == embedding count (no orphans either way)
  - embedding batch wall-clock on a realistically sized document

The link's external HTTP fetch is mocked (links._MOCK) for determinism — the
TinyFish call is a thin requests.post mirroring tinyfish_service; everything
downstream of the fetch (orchestration, chunk, embed of link text) is real.
"""
import io
import os
import sqlite3
import time

import pytest

# Import first — these trigger load_dotenv(), populating os.environ from .env so
# the key check below sees a .env-only key (not just shell env).
from backend.database.schema import ALL_TABLES, MIGRATIONS
from backend.services.feed_v2 import db as v2db
from backend.services.feed_v2.schema import run_v2_migrations
from backend.services.feed_v2.ingestion import materials, links, images
from backend.services.feed_v2.llm import provider

_HAS_KEY = bool(os.getenv("GEMINI_API_KEYS") or os.getenv("GEMINI_API_KEY"))
pytestmark = [pytest.mark.integration,
              pytest.mark.skipif(not _HAS_KEY, reason="no Gemini API key")]


def _build_db(path):
    conn = sqlite3.connect(path)
    for stmt in ALL_TABLES:
        conn.execute(stmt)
    for m in MIGRATIONS:
        try:
            conn.execute(m) if not isinstance(m, (list, tuple)) else [conn.execute(s) for s in m]
        except sqlite3.OperationalError as e:
            if not any(p in str(e).lower() for p in ("already exists", "duplicate column", "no such column")):
                raise
    run_v2_migrations(conn)
    conn.execute("INSERT OR IGNORE INTO users(user_id,email,name,hashed_pw) VALUES('u1','u1@t.com','u','x')")
    conn.execute("INSERT OR IGNORE INTO v2_projects(project_id,user_id,name) VALUES('proj','u1','P')")
    conn.commit()
    conn.close()


@pytest.fixture
def db(tmp_path, monkeypatch):
    path = str(tmp_path / "ingest.db")
    _build_db(path)
    monkeypatch.setattr(v2db, "DB_PATH", path)
    monkeypatch.setattr(links, "_MOCK", True)  # deterministic link fetch
    return path


# ── fixture builders (real files) ─────────────────────────────────────────────
def _png(text="HELLO 42"):
    from PIL import Image, ImageDraw
    img = Image.new("RGB", (240, 120), "white")
    ImageDraw.Draw(img).text((20, 50), text, fill="black")
    buf = io.BytesIO(); img.save(buf, format="PNG"); return buf.getvalue()


def _pdf_with_figure():
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.utils import ImageReader
    from PIL import Image
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    c.drawString(72, 740, "Introduction to Widgets. " * 6)
    c.drawString(72, 720, "This page has real digital text well over the scanned-PDF threshold. " * 2)
    red = Image.new("RGB", (150, 150), "red")
    c.drawImage(ImageReader(red), 72, 420, width=180, height=180)
    c.drawString(72, 400, "Figure 1: a red square widget diagram")
    c.showPage()
    c.drawString(72, 740, "Chapter 2 discusses widget assembly in extended detail. " * 8)
    c.showPage()
    c.save()
    return buf.getvalue()


def _docx_with_image():
    from docx import Document
    doc = Document()
    doc.add_heading("Overview", level=1)
    doc.add_paragraph("Widgets are small components. " * 20)
    doc.add_heading("Details", level=1)
    doc.add_paragraph("Assembly requires care. " * 20)
    doc.add_picture(io.BytesIO(_png("DOCX FIG")))
    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()


def _big_text(n_paras=60):
    return "\n\n".join(
        f"Paragraph {i}: " + ("widget assembly detail and reasoning " * 12) for i in range(n_paras)
    ).encode("utf-8")


def _chunk_counts(material_id):
    with v2db.get_connection() as conn:
        c = conn.execute("SELECT COUNT(*) FROM v2_material_chunks WHERE material_id=?", (material_id,)).fetchone()[0]
        v = conn.execute("SELECT COUNT(*) FROM v2_material_chunks_vec WHERE material_id=?", (material_id,)).fetchone()[0]
    return c, v


def test_mixed_upload_set_lands_with_figure_and_matched_counts(db, capsys):
    pdf = materials.ingest(u := "u1", "proj", file_bytes=_pdf_with_figure(), filename="widgets.pdf")
    docx = materials.ingest("u1", "proj", file_bytes=_docx_with_image(), filename="widgets.docx")
    img = materials.ingest("u1", "proj", file_bytes=_png(), filename="photo.png")
    link = materials.ingest("u1", "proj", url="https://example.com/widgets")

    # 1. All four land with correct type + status done.
    with v2db.get_connection() as conn:
        rows = {r["material_id"]: r for r in conn.execute(
            "SELECT material_id, type, extraction_status FROM v2_materials WHERE project_id='proj'").fetchall()}
    assert len(rows) == 4
    assert {r["type"] for r in rows.values()} == {"document", "image", "link"}
    assert all(r["extraction_status"] == "done" for r in rows.values()), rows

    # 2. At least one figure from the PDF with a real page number + caption.
    with v2db.get_connection() as conn:
        figs = conn.execute(
            "SELECT page_no, caption FROM v2_material_figures WHERE material_id=?", (pdf["material_id"],)).fetchall()
    assert figs, "no figure extracted from the PDF"
    assert any(f["page_no"] is not None for f in figs)
    src = pdf["figures"][0]["caption_source"]
    with capsys.disabled():
        print(f"\nPDF figure: page_no={figs[0]['page_no']} caption={figs[0]['caption']!r} source={src}")

    # 3. chunk count == embedding count for every text-bearing material (no orphans).
    for mid, r in rows.items():
        c, v = _chunk_counts(mid)
        assert c == v, f"{r['type']} {mid}: chunks={c} embeddings={v}"
        if r["type"] in ("document", "link"):
            assert c > 0, f"{r['type']} produced no chunks"

    with capsys.disabled():
        print("chunk==embedding counts:",
              {r["type"]: _chunk_counts(mid)[0] for mid, r in rows.items()})


def test_image_ingestor_fallback_leg_serves_through_routing_table(monkeypatch, capsys):
    """FLAG 3: prove a photo upload's output comes through image_ingestor's
    primary/fallback ROUTING TABLE, not some bypass path. Force the primary leg
    (gemini-3-flash-preview) to fail; confirm the fallback leg
    (gemini-3.1-flash-lite) is the one that actually serves a real request —
    same proof style as Phase 3 key rotation. describe_image -> call_agent
    ('image_ingestor') is the single entry point (images.py:45)."""
    primary_id = provider.MODEL_REGISTRY["gemini-3-flash-preview"][1]
    fallback_id = provider.MODEL_REGISTRY["gemini-3.1-flash-lite"][1]
    real_google = provider._call_google
    attempted: list[str] = []
    served: list[str] = []

    def fake_google(api_model_id, messages, system, schema, key, images=None):
        if api_model_id == primary_id:
            attempted.append(api_model_id)
            raise RuntimeError("simulated primary outage")  # non-429 -> immediate leg abort, no cooldown pollution
        served.append(api_model_id)
        return real_google(api_model_id, messages, system, schema, key, images)

    monkeypatch.setitem(provider._SDK_FOR_PROVIDER, "google", fake_google)
    monkeypatch.setattr(provider.call_logger, "write_call_row", lambda **k: None)  # isolate routing; no DB

    res = images.describe_image(_png("ROUTE 7"), ".png",
                                meta={"call_type": "p4b_route_test", "surface": "feed_v2"})

    assert res.error is None, res.error
    assert attempted and all(a == primary_id for a in attempted), attempted  # primary WAS tried
    assert served == [fallback_id], served  # fallback served exactly once; primary NOT re-hit as the server
    assert res.description or res.ocr_text  # a real Gemini response came back through the fallback leg
    with capsys.disabled():
        print(f"\nimage_ingestor routing: primary {primary_id} failed -> fallback {served[0]} served "
              f"(desc_len={len(res.description)}, ocr_len={len(res.ocr_text)})")


def test_embedding_batch_walltime_on_realistic_doc(db, capsys):
    big = _big_text(60)  # ~60 paragraphs -> many 800-char chunks -> multiple batches
    t0 = time.monotonic()
    res = materials.ingest("u1", "proj", file_bytes=big, filename="manual.txt")
    elapsed = time.monotonic() - t0
    assert res["status"] == "done"
    c, v = _chunk_counts(res["material_id"])
    assert c == v > 0
    with capsys.disabled():
        print(f"\nREALISTIC DOC: {len(big)} bytes -> {c} chunks, batched embed+store "
              f"wall-clock {elapsed:.2f}s")
