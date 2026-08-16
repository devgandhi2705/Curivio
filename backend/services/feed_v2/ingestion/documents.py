"""
Feed v2 document extraction — full text + SHA-256 + structure (TOC/headings).

Reuses the SAME libraries the chat path's document_extraction_service wraps
(pypdf, python-docx) DIRECTLY — importing that service would cross the isolation
boundary. The chat wrapper's scanned-PDF guard (<100 chars/page → clear error,
no fake OCR) is replicated here; pptx is intentionally out of scope this phase.

extract(file_bytes, filename, ext) -> ExtractResult with exactly one of
(text populated + error None) or (text None + error populated).
"""
from __future__ import annotations

import hashlib
import io
import logging
from dataclasses import dataclass, field

import pypdf
from docx import Document

logger = logging.getLogger(__name__)

# pptx dropped this phase (python-pptx not installed — see Phase 4 decision).
DOCUMENT_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}
_PLAIN_EXTENSIONS = {".txt", ".md"}

# Same threshold the chat path uses: a real digital PDF page runs into the
# hundreds of chars; a scanned/image-only page extracts ~0 (pypdf has no OCR).
_MIN_CHARS_PER_PAGE = 100


@dataclass
class ExtractResult:
    text: str | None = None
    error: str | None = None
    sha256: str = ""
    byte_size: int = 0
    page_count: int | None = None
    structure: list[dict] = field(default_factory=list)  # [{title, level, page?}]
    # Phase 8b: per-page text for PDFs so chunking can carry a real, exact page_no
    # into each chunk (documents that have no page concept — docx/txt/md — leave this
    # None and their chunks stay page_no=NULL). Each entry: {"page_no": int (1-based),
    # "text": str}. `text` above stays the flat concatenation for callers that want
    # the whole document; `pages` is the structured view chunking uses.
    pages: list[dict] | None = None

    @property
    def has_structure(self) -> bool:
        return bool(self.structure)

    @property
    def section_count(self) -> int:
        return len(self.structure)


def sha256_hex(file_bytes: bytes) -> str:
    return hashlib.sha256(file_bytes).hexdigest()


def _pdf_outline(reader: pypdf.PdfReader) -> list[dict]:
    """Flatten PDF bookmarks/outline into ordered {title, level, page} entries.
    A real TOC/chapter sequence lives here; many PDFs have none (empty list)."""
    out: list[dict] = []

    def walk(items, level=0):
        for it in items:
            if isinstance(it, list):
                walk(it, level + 1)
                continue
            title = getattr(it, "title", None)
            if not title:
                continue
            entry = {"title": str(title).strip(), "level": level}
            try:
                entry["page"] = reader.get_destination_page_number(it)
            except Exception:
                pass
            out.append(entry)

    try:
        walk(reader.outline)
    except Exception:
        logger.debug("pdf outline walk failed", exc_info=True)
    return out


def _extract_pdf(file_bytes: bytes, filename: str) -> ExtractResult:
    r = ExtractResult(sha256=sha256_hex(file_bytes), byte_size=len(file_bytes))
    try:
        reader = pypdf.PdfReader(io.BytesIO(file_bytes))
    except Exception as exc:
        r.error = f"Couldn't read '{filename}' as a PDF: {exc}"
        return r
    if not reader.pages:
        r.error = f"'{filename}' has no pages."
        return r
    page_texts = [(p.extract_text() or "") for p in reader.pages]
    r.page_count = len(page_texts)
    if sum(len(t) for t in page_texts) / len(page_texts) < _MIN_CHARS_PER_PAGE:
        r.error = (f"Couldn't extract text from '{filename}' — looks like a scanned/"
                   f"image-only PDF. OCR isn't available.")
        return r
    r.text = "\n\n".join(page_texts)                     # flat view (unchanged)
    # Per-page view for page-aware chunking. 1-based page_no is what a citation shows
    # a learner ("page 4" = the 4th page), unlike pypdf's 0-based internal indices.
    r.pages = [{"page_no": i + 1, "text": t} for i, t in enumerate(page_texts)]
    r.structure = _pdf_outline(reader)                   # reads reader.outline, NOT r.text
    return r


def _docx_headings(doc: Document) -> list[dict]:
    out: list[dict] = []
    for p in doc.paragraphs:
        style = (p.style.name if p.style else "") or ""
        if style.startswith("Heading") and p.text.strip():
            level = "".join(ch for ch in style if ch.isdigit())
            out.append({"title": p.text.strip(), "level": int(level) if level else 1})
    return out


def _extract_docx(file_bytes: bytes, filename: str) -> ExtractResult:
    r = ExtractResult(sha256=sha256_hex(file_bytes), byte_size=len(file_bytes))
    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as exc:
        r.error = f"Couldn't read '{filename}' as a .docx: {exc}"
        return r
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    text = "\n".join(parts)
    if not text.strip():
        r.error = f"'{filename}' has no extractable text."
        return r
    r.text = text
    r.structure = _docx_headings(doc)
    return r


def _md_headings(text: str) -> list[dict]:
    out: list[dict] = []
    for line in text.splitlines():
        s = line.lstrip()
        if s.startswith("#"):
            level = len(s) - len(s.lstrip("#"))
            title = s[level:].strip()
            if title:
                out.append({"title": title, "level": level})
    return out


def _extract_plain(file_bytes: bytes, filename: str, ext: str) -> ExtractResult:
    r = ExtractResult(sha256=sha256_hex(file_bytes), byte_size=len(file_bytes))
    try:
        text = file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        try:
            text = file_bytes.decode("latin-1")
        except Exception as exc:
            r.error = f"Couldn't decode '{filename}' as text: {exc}"
            return r
    if not text.strip():
        r.error = f"'{filename}' is empty."
        return r
    r.text = text
    if ext == ".md":
        r.structure = _md_headings(text)
    return r


def extract(file_bytes: bytes, filename: str, ext: str) -> ExtractResult:
    ext = ext.lower()
    if ext == ".pdf":
        return _extract_pdf(file_bytes, filename)
    if ext == ".docx":
        return _extract_docx(file_bytes, filename)
    if ext in _PLAIN_EXTENSIONS:
        return _extract_plain(file_bytes, filename, ext)
    r = ExtractResult(sha256=sha256_hex(file_bytes), byte_size=len(file_bytes))
    r.error = f"Unsupported file type: {ext or 'unknown'}"
    return r
