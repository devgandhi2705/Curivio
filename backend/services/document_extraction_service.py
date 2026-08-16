"""
Chat-R6a — direct text extraction for uploaded documents (pdf/docx/csv/text/code).

Images stay entirely on the existing native Gemini vision path (chat_agent.py's
Chat-5 has_attachments gate) — this module is never invoked for image uploads.

Scanned/image-only PDFs are NOT OCR'd (R6b, separate/unresolved phase) — pypdf's
extracted text is measured against a real <100 chars/page threshold (live-verified:
a hand-built PDF with real text content extracted to 73 chars over 2 short lines;
a page with no text content stream extracted to exactly 0 chars) and a clear,
specific error is returned instead of a silent wrong answer or a fake OCR attempt.

Routes by file extension, not content_type: browsers send unreliable/missing MIME
types for code files (live-verified via mimetypes.guess_type as a proxy — .ts
guesses as a video MIME, .tsx/.jsx/.yaml/.yml/.java/.go/.rs/.cpp all guess None).
docx/pdf have consistent MIME types across browsers so main.py's image-vs-document
routing can still use content_type; extraction itself uses the extension.

Public API
----------
DOCUMENT_EXTENSIONS: set[str]   every extension this module can extract
extract_document_text(file_bytes, filename, ext) -> (text, error, error_type)
    Exactly one of (text) or (error, error_type) is populated. error_type is
    one of the ERROR_* constants below — a stable machine code for each of
    the 8 distinct failure reasons this module can produce (Phase P), so
    /chat/upload's llm_call_log row carries the real reason instead of a
    caller re-deriving it by string-matching `error`.
"""
from __future__ import annotations

import io
import logging

import pypdf
from docx import Document

logger = logging.getLogger(__name__)

# Real digital PDF pages run into the hundreds of chars/page; a scanned/image-only
# page extracts to 0 chars (pypdf has nothing to read — no OCR). Threshold sits
# comfortably below real prose while staying well above the scanned-page floor.
_MIN_CHARS_PER_PAGE = 100

# Phase P — one code per distinct failure reason, logged verbatim as
# llm_call_log.error_type for the 'extract' step of a chat_upload row.
ERROR_PDF_UNREADABLE        = "pdf_unreadable"
ERROR_PDF_NO_PAGES          = "pdf_no_pages"
ERROR_PDF_SCANNED           = "pdf_scanned"
ERROR_DOCX_UNREADABLE       = "docx_unreadable"
ERROR_DOCX_NO_TEXT          = "docx_no_text"
ERROR_TEXT_UNDECODABLE      = "text_undecodable"
ERROR_TEXT_EMPTY            = "text_empty"
ERROR_UNSUPPORTED_EXTENSION = "unsupported_extension"

_PLAIN_TEXT_EXTENSIONS = {
    ".txt", ".md", ".csv",
    ".py", ".js", ".jsx", ".ts", ".tsx", ".java", ".go", ".rs", ".c", ".cpp",
    ".h", ".hpp", ".cs", ".rb", ".php", ".sh", ".sql", ".json", ".yaml", ".yml",
    ".html", ".css", ".xml",
}

DOCUMENT_EXTENSIONS = _PLAIN_TEXT_EXTENSIONS | {".pdf", ".docx"}


def _extract_pdf(file_bytes: bytes, filename: str) -> tuple[str | None, str | None, str | None]:
    try:
        reader = pypdf.PdfReader(io.BytesIO(file_bytes))
    except Exception as exc:
        return None, f"Couldn't read '{filename}' as a PDF: {exc}", ERROR_PDF_UNREADABLE

    if not reader.pages:
        return None, f"'{filename}' has no pages.", ERROR_PDF_NO_PAGES

    page_texts = [(p.extract_text() or "") for p in reader.pages]
    total_chars = sum(len(t) for t in page_texts)
    if total_chars / len(page_texts) < _MIN_CHARS_PER_PAGE:
        return None, (
            f"Couldn't extract text from '{filename}' — this looks like a "
            f"scanned or image-only PDF. OCR support isn't available yet "
            f"(coming in a future update)."
        ), ERROR_PDF_SCANNED
    return "\n\n".join(page_texts), None, None


def _extract_docx(file_bytes: bytes, filename: str) -> tuple[str | None, str | None, str | None]:
    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as exc:
        return None, f"Couldn't read '{filename}' as a .docx file: {exc}", ERROR_DOCX_UNREADABLE

    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    text = "\n".join(parts)
    if not text.strip():
        return None, f"'{filename}' has no extractable text.", ERROR_DOCX_NO_TEXT
    return text, None, None


def _extract_plain(file_bytes: bytes, filename: str) -> tuple[str | None, str | None, str | None]:
    try:
        text = file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        try:
            text = file_bytes.decode("latin-1")
        except Exception as exc:
            return None, f"Couldn't decode '{filename}' as text: {exc}", ERROR_TEXT_UNDECODABLE
    if not text.strip():
        return None, f"'{filename}' is empty.", ERROR_TEXT_EMPTY
    return text, None, None


def extract_document_text(file_bytes: bytes, filename: str, ext: str) -> tuple[str | None, str | None, str | None]:
    """Returns (text, error, error_type) — exactly one of (text) or
    (error, error_type) is populated."""
    ext = ext.lower()
    if ext == ".pdf":
        return _extract_pdf(file_bytes, filename)
    if ext == ".docx":
        return _extract_docx(file_bytes, filename)
    if ext in _PLAIN_TEXT_EXTENSIONS:
        return _extract_plain(file_bytes, filename)
    return None, f"Unsupported file type: {ext or 'unknown'}", ERROR_UNSUPPORTED_EXTENSION
